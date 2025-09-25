"""Regenerate valid DOCX files from the existing plain-text transcript files.

Current state: transcript_PATIENT_*.docx files were created as plain UTF-8 text but
named with a .docx extension, so they are not valid Word (OpenXML) packages.

This script:
 1. Reads each transcript_*.docx as plain text (since they are not true docx yet)
 2. Parses line-by-line, removing simple markdown bold markers (**)
 3. Rebuilds a proper .docx using python-docx
 4. Applies Arial 12pt as requested
 5. Bold-styles lines that originally contained ** markers (questions, sections)

Safety:
 - Skips any file that already appears to be a valid docx (ZIP magic PK) unless --force specified
 - Creates a backup copy with .orig.txt if converting an invalid placeholder
"""

from __future__ import annotations
import argparse
import sys
from pathlib import Path
from typing import Iterable
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).parent

def is_valid_docx(path: Path) -> bool:
    try:
        with path.open('rb') as f:
            sig = f.read(2)
        return sig == b'PK'
    except OSError:
        return False

def iter_transcript_files(pattern: str = 'transcript_PATIENT_*.docx') -> Iterable[Path]:
    return sorted(ROOT.glob(pattern))

def apply_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Arial'
        # Ensure compatibility with Word's East Asia font mapping
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(12)

def convert_plaintext_to_docx(src: Path, force: bool = False, verbose: bool = True) -> bool:
    if is_valid_docx(src) and not force:
        if verbose:
            print(f"SKIP (already valid): {src.name}")
        return False
    try:
        text = src.read_text(encoding='utf-8', errors='replace')
    except UnicodeDecodeError:
        if verbose:
            print(f"ERROR: Cannot read as text: {src}")
        return False

    backup = src.with_suffix(src.suffix + '.orig.txt')
    if not backup.exists():
        backup.write_text(text, encoding='utf-8')

    doc = Document()
    # Set default (Normal) style to Arial 12
    # Access normal style font (guard if missing)
    try:
        style = doc.styles['Normal']
        style.font.name = 'Arial'  # type: ignore[attr-defined]
        style.font.size = Pt(12)   # type: ignore[attr-defined]
    except Exception:
        pass

    for raw_line in text.splitlines():
        line = raw_line.rstrip('\r')
        if not line.strip():
            doc.add_paragraph("")
            continue
        had_bold_marker = '**' in line
        clean_line = line.replace('**', '')
        p = doc.add_paragraph()
        run = p.add_run(clean_line)
        if had_bold_marker:
            run.bold = True
        apply_font(p)

    # Save over original path (now becomes a valid docx)
    doc.save(str(src))
    if verbose:
        print(f"Converted -> {src.name}")
    return True

def main(argv=None):
    parser = argparse.ArgumentParser(description="Regenerate valid DOCX transcripts")
    parser.add_argument('--force', action='store_true', help='Rebuild even if file already valid')
    parser.add_argument('--pattern', default='transcript_PATIENT_*.docx')
    args = parser.parse_args(argv)

    files = list(iter_transcript_files(args.pattern))
    if not files:
        print("No transcript files found matching pattern", file=sys.stderr)
        return 1

    converted = 0
    for f in files:
        if convert_plaintext_to_docx(f, force=args.force):
            converted += 1
    print(f"Done. {converted} file(s) converted.")
    return 0

if __name__ == '__main__':  # pragma: no cover
    raise SystemExit(main())
